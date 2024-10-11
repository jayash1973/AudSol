import streamlit as st
import requests
import json
import os
import pyaudio as pyaudio
import base64
from collections import defaultdict
from io import BytesIO
import speech_recognition as sr
from elevenlabs import play, stream, save
from elevenlabs.client import ElevenLabs
from elevenlabs import Voice, VoiceSettings
from together import Together
import tempfile
from audio_recorder_streamlit import audio_recorder
from PIL import Image
import random
import string
import av
import time
import io
import queue
from docx import Document
from fpdf import FPDF
import pytesseract
import plotly.express as px
import pandas as pd
from pdf2image import convert_from_path
import PyPDF2
import docx2txt
from streamlit_option_menu import option_menu
import nltk
from nltk.tokenize import sent_tokenize
import networkx as nx
import speech_recognition as sr
from streamlit_webrtc import webrtc_streamer, WebRtcMode, RTCConfiguration
import matplotlib.pyplot as plt
import plotly.graph_objects as go
from plotly.subplots import make_subplots

nltk.download('punkt')

from streamlit_agraph import agraph, Node, Edge, Config
import pandas as pd
from wordcloud import WordCloud
import numpy as np
import threading
import pydub
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import math
import wave
import re
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# API Keys
XI_API_KEY = "sk_6311dfd264e562763d5c2dd93acad19a020e91efdb281bfa"
TOGETHER_API_KEY = "291a366758796aaf86fe851b1b0ce2278ef75c750930e9f5d949e6b38bd208de"

# Initialize clients
elevenlabs_client = ElevenLabs(api_key=XI_API_KEY)
together_client = Together(api_key=TOGETHER_API_KEY)


# Predefined characters (expanded)
predefined_characters = {
    "Naruto Uzumaki": {
        "description": "Protagonist of the Naruto series. An optimistic and determined ninja with the goal of becoming Hokage.",
        "voice_id": "ErXwobaYiN019PkySvjV",
        "image": "https://i.pinimg.com/originals/2e/f7/0d/2ef70d5217b530dfb766a45d9babbb5e.jpg"
        },
    "Sherlock Holmes": {
        "description": "Brilliant detective known for his logical reasoning and observational skills.",
        "voice_id": "VR6AewLTigWG4xSOukaG",
        "image": "https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcSFNrsQ14wE7vKH_46oN-tCF3YwGyjo9fbLMuTpvm3ENf_10JcTHqoIyxkl_EDgpmGnEXs&usqp=CAU"
    },
    "Elizabeth Bennet": {
        "description": "Protagonist of Pride and Prejudice, known for her intelligence and wit.",
        "voice_id": "EXAVITQu4vr4xnSDxMaL",
        "image": "https://www.indiependent.co.uk/wp-content/uploads/2015/08/elizabeth-bennet.jpg"
    }
}

# Initialize session state
if "characters" not in st.session_state or not st.session_state.characters:
    st.session_state.characters = predefined_characters.copy()
if "current_character" not in st.session_state:
    st.session_state.current_character = None
if "messages" not in st.session_state:
    st.session_state.messages = []
if "book_details" not in st.session_state:
    st.session_state.book_details = {}

# Fiction genres
fiction_genres = [
    "Fantasy", "Science Fiction", "Mystery", "Thriller", "Romance", "Historical Fiction",
    "Horror", "Adventure", "Contemporary Fiction", "Dystopian", "Young Adult"
]

# Non-fiction genres
non_fiction_genres = [
    "Biography", "Autobiography", "Memoir", "Self-help", "History", "Science",
    "Philosophy", "Psychology", "Business", "Travel", "True Crime"
]

def create_character(name, description, voice_file, image_file=None):
    voice = elevenlabs_client.clone(
        name=name,
        description=description,
        files=[voice_file]
    )
    
    image_path = None
    if image_file:
        characters_dir = "characters"
        os.makedirs(characters_dir, exist_ok=True)
        image_path = os.path.join(characters_dir, f"{name.lower().replace(' ', '_')}.jpg")
        with Image.open(image_file) as img:
            img.save(image_path)
    
    st.session_state.characters[name] = {
        "description": description,
        "voice_id": voice.voice_id,
        "image": image_path
    }
    return voice

def generate_ai_response(character, user_message, language):
    response = together_client.chat.completions.create(
        model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
        messages=[
            {"role": "system", "content": f"You are {character}. {st.session_state.characters[character]['description']} Respond in {language}."},
            {"role": "user", "content": user_message}
        ],
        max_tokens=1024,
        temperature=0.7,
        top_p=0.7,
        top_k=50,
        repetition_penalty=1,
        stop=["<|eot_id|>", "<|eom_id|>"]
    )
    return response.choices[0].message.content

def text_to_speech(text, voice_id):
    audio = elevenlabs_client.generate(
        text=text,
        voice=Voice(voice_id=voice_id),
        model="eleven_multilingual_v2"
    )
    return audio

def speech_to_text(audio_bytes):
    recognizer = sr.Recognizer()
    with sr.AudioFile(BytesIO(audio_bytes)) as source:
        audio = recognizer.record(source)
    try:
        text = recognizer.recognize_google(audio)
        return text
    except sr.UnknownValueError:
        return "Sorry, I couldn't understand the audio."
    except sr.RequestError:
        return "Sorry, there was an error processing the audio."

def process_audio(frame):
    sound = frame.to_ndarray()
    sound = sound.astype(np.int16)
    return sound.tobytes()

def split_text(text, chunk_size=5000):
    """Split the text into chunks of approximately chunk_size characters."""
    chunks = []
    current_chunk = []
    current_size = 0
    for sentence in nltk.sent_tokenize(text):
        sentence_size = len(sentence)
        if current_size + sentence_size > chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            current_chunk = []
            current_size = 0
        current_chunk.append(sentence)
        current_size += sentence_size
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    return chunks

def generate_audio_chunk(text, voice_id, headers):
    """Generate audio for a single chunk of text."""
    audio_data = {
        "text": text,
        "model_id": "eleven_monolingual_v1",
        "voice_settings": {
            "stability": 0.5,
            "similarity_boost": 0.5
        }
    }
    response = requests.post(f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}", json=audio_data, headers=headers)
    if response.status_code == 200:
        return response.content
    else:
        raise Exception(f"Error generating audio: {response.text}")

def convert_to_audiobook(text, language, voice_id):
    audio = elevenlabs_client.generate(
        text=text,
        voice=Voice(voice_id=voice_id),
        model="eleven_multilingual_v2"
    )
    return audio
def generate_book_chapter(book_details, chapter_index):
    chapter = book_details['chapters'][chapter_index]
    chapter_name = chapter['name']

    if book_details['genre_type'] == 'Fiction':
        chapter_prompt = f"Write chapter {chapter_index + 1} titled '{chapter_name}' for a {book_details['genre']} book. The book is about: {book_details['description']}. This chapter should include the following scenes: {', '.join(chapter['scenes'])}. Maintain consistency with the overall story and characters."
        
        chapter_content = ""
        for scene_index, scene in enumerate(chapter['scenes']):
            scene_prompt = f"Write scene {scene_index + 1} for chapter {chapter_index + 1} titled '{chapter_name}'. The scene should include the following details: {scene}"
            scene_text = together_client.chat.completions.create(
                model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
                messages=[
                    {"role": "system", "content": f"You are an expert {book_details['genre']} writer."},
                    {"role": "user", "content": scene_prompt}
                ],
                max_tokens=6048,
                temperature=0.7
            )
            chapter_content += f"## Scene {scene_index + 1}\n\n{scene_text.choices[0].message.content}\n\n"
    else:
        chapter_prompt = f"Write chapter {chapter_index + 1} titled '{chapter_name}' for a {book_details['genre']} non-fiction book. The book is about: {book_details['description']}. This chapter should cover the following parts: {', '.join(chapter['parts'])}. Ensure the content is informative and well-structured."
        
        chapter_content = ""
        for part_index, part in enumerate(chapter['parts']):
            part_prompt = f"Write part {part_index + 1} for chapter {chapter_index + 1} titled '{chapter_name}'. The part should cover the following details: {part}"
            part_text = together_client.chat.completions.create(
                model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
                messages=[
                    {"role": "system", "content": f"You are an expert {book_details['genre']} writer."},
                    {"role": "user", "content": part_prompt}
                ],
                max_tokens=6048,
                temperature=0.7
            )
            chapter_content += f"## Part {part_index + 1}\n\n{part_text.choices[0].message.content}\n\n"
    
    return chapter_content
def generate_book_introduction(book_details):
    intro_prompt = f"Write an engaging introduction for a {book_details['genre']} {'fiction' if book_details['genre_type'] == 'Fiction' else 'non-fiction'} book titled '{book_details['title']}' by {book_details['author']}. The book is about: {book_details['description']}."
    
    intro_content = together_client.chat.completions.create(
        model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
        messages=[
            {"role": "system", "content": "You are an expert book introduction writer."},
            {"role": "user", "content": intro_prompt}
        ],
        max_tokens=1024,
        temperature=0.7
    )
    return intro_content.choices[0].message.content

def generate_table_of_contents(book_details):
    toc = f"# Table of Contents\n\nIntroduction\n\n"
    for i, chapter in enumerate(book_details['chapters'], 1):
        toc += f"{i}. {chapter['name']}\n"
    return toc

def save_book_formats(content, title):
    # Save as TXT
    with open(f"{title}.txt", "w", encoding="utf-8") as f:
        f.write(content)

    # Save as DOCX
    doc = Document()
    doc.add_heading(title, 0)
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)
    doc.save(f"{title}.docx")

    # Save as PDF using reportlab
    pdf_filename = f"{title}.pdf"
    doc = SimpleDocTemplate(pdf_filename, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = []

    for paragraph in content.split('\n'):
        p = Paragraph(paragraph, styles['Normal'])
        flowables.append(p)

    doc.build(flowables)

def extract_text_from_file(file):
    file_extension = os.path.splitext(file.name)[1].lower()
    
    if file_extension == '.pdf':
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    elif file_extension == '.txt':
        text = file.getvalue().decode('utf-8')
    elif file_extension in ['.doc', '.docx']:
        text = docx2txt.process(file)
    else:
        raise ValueError("Unsupported file format")
    
    return text


# World-Building Assistant
def world_building_assistant():
    st.header("World-Building Assistant")
    st.write("Develop and maintain consistent rules, cultures, and environments for your fictional world with AI assistance.")

    if 'world_elements' not in st.session_state:
        st.session_state.world_elements = {}

    # Input for world elements
    element_category = st.selectbox("Element Category", ["Geography", "Culture", "Magic System", "Technology", "History", "Politics"])

    # Custom inputs for each category
    if element_category == "Geography":
        element_name = st.text_input("Location Name")
        climate = st.selectbox("Climate", ["Tropical", "Temperate", "Arctic", "Desert", "Mediterranean"])
        terrain = st.multiselect("Terrain Features", ["Mountains", "Forests", "Rivers", "Oceans", "Plains", "Islands"])
        prompt = f"Create a detailed description for a {climate} {', '.join(terrain)} region named {element_name}."

    elif element_category == "Culture":
        element_name = st.text_input("Culture Name")
        social_structure = st.selectbox("Social Structure", ["Hierarchical", "Egalitarian", "Clan-based", "Caste System"])
        values = st.multiselect("Core Values", ["Honor", "Knowledge", "Nature", "Technology", "Spirituality", "Warfare"])
        prompt = f"Describe the {social_structure} culture of {element_name}, emphasizing their focus on {', '.join(values)}."

    elif element_category == "Magic System":
        element_name = st.text_input("Magic System Name")
        source = st.selectbox("Source of Magic", ["Natural Elements", "Divine", "Inner Energy", "Artifacts", "Otherworldly"])
        limitations = st.multiselect("Limitations", ["Physical toll", "Rare resources", "Specific bloodline", "Years of study", "Unpredictable"])
        prompt = f"Detail the {source}-based magic system called {element_name}, including its {', '.join(limitations)} as limitations."

    elif element_category == "Technology":
        element_name = st.text_input("Technology Name")
        tech_level = st.slider("Technology Level", 1, 10, 5)
        focus = st.multiselect("Technological Focus", ["Energy", "Transportation", "Communication", "Warfare", "Medicine", "AI"])
        prompt = f"Describe the level {tech_level} technology {element_name}, focusing on advancements in {', '.join(focus)}."

    elif element_category == "History":
        element_name = st.text_input("Historical Event/Era Name")
        time_frame = st.selectbox("Time Frame", ["Ancient", "Medieval", "Renaissance", "Industrial", "Modern", "Futuristic"])
        event_type = st.selectbox("Event Type", ["War", "Discovery", "Cultural Revolution", "Natural Disaster", "Technological Breakthrough"])
        prompt = f"Narrate the {time_frame} {event_type} known as {element_name} and its impact on the world."

    elif element_category == "Politics":
        element_name = st.text_input("Political System/Faction Name")
        gov_type = st.selectbox("Government Type", ["Monarchy", "Democracy", "Oligarchy", "Theocracy", "Anarchy"])
        key_issues = st.multiselect("Key Political Issues", ["Resource scarcity", "Technological regulation", "Civil rights", "Environmental concerns", "Foreign relations"])
        prompt = f"Explain the {gov_type} political system of {element_name}, addressing their stance on {', '.join(key_issues)}."

    if st.button("Generate World Element"):
        if element_name:
            with st.spinner("Generating world element..."):
                response = together_client.chat.completions.create(
                    model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
                    messages=[
                        {"role": "system", "content": "You are an expert world-building assistant for fiction writers."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=5000,
                    temperature=0.7
                )
                element_description = response.choices[0].message.content

            if element_category not in st.session_state.world_elements:
                st.session_state.world_elements[element_category] = {}
            st.session_state.world_elements[element_category][element_name] = element_description
            st.success(f"{element_name} added to {element_category} successfully!")

    # Display world elements
    if st.session_state.world_elements:
        st.subheader("Your World")
        for category, elements in st.session_state.world_elements.items():
            st.write(f"**{category}**")
            for name, description in elements.items():
                with st.expander(name):
                    st.write(description)
                    if st.button(f"Regenerate {name}", key=f"regen_{category}_{name}"):
                        with st.spinner(f"Regenerating {name}..."):
                            response = together_client.chat.completions.create(
                                model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
                                messages=[
                                    {"role": "system", "content": "You are an expert world-building assistant for fiction writers."},
                                    {"role": "user", "content": f"Rewrite and improve the following world element description for {category}: {description}"}
                                ],
                                max_tokens=500,
                                temperature=0.7
                            )
                            new_description = response.choices[0].message.content
                            st.session_state.world_elements[category][name] = new_description
                            st.experimental_rerun()

    # AI-powered world consistency check
    if st.button("Check World Consistency"):
        if st.session_state.world_elements:
            all_elements = "\n".join([f"{cat}: {name} - {desc}" for cat, elements in st.session_state.world_elements.items() for name, desc in elements.items()])
            with st.spinner("Analyzing world consistency..."):
                response = together_client.chat.completions.create(
                    model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
                    messages=[
                        {"role": "system", "content": "You are an expert world-building consultant for fiction writers."},
                        {"role": "user", "content": f"Analyze the following world elements for consistency and provide suggestions for improvement:\n\n{all_elements}"}
                    ],
                    max_tokens=1000,
                    temperature=0.7
                )
                consistency_analysis = response.choices[0].message.content
            st.subheader("World Consistency Analysis")
            st.write(consistency_analysis)
        else:
            st.warning("Add some world elements before checking consistency.")

    # Export world-building data
    if st.session_state.world_elements:
        if st.button("Export World Data"):
            export_data = {category: {name: desc for name, desc in elements.items()} for category, elements in st.session_state.world_elements.items()}
            st.download_button(
                label="Download World Data as JSON",
                data=json.dumps(export_data, indent=2),
                file_name="world_building_data.json",
                mime="application/json"
            )

def get_audio_data():
    CHUNK = 1024
    FORMAT = pyaudio.paInt16
    CHANNELS = 1
    RATE = 44100
    RECORD_SECONDS = 0.1

    p = pyaudio.PyAudio()

    stream = p.open(format=FORMAT,
                    channels=CHANNELS,
                    rate=RATE,
                    input=True,
                    frames_per_buffer=CHUNK)

    frames = []

    for i in range(0, int(RATE / CHUNK * RECORD_SECONDS)):
        data = stream.read(CHUNK)
        frames.append(np.frombuffer(data, dtype=np.int16))

    stream.stop_stream()
    stream.close()
    p.terminate()

    return np.concatenate(frames)

def update_audio_chart(chart, audio_data):
    chart.plotly_chart(go.Figure(data=go.Scatter(y=audio_data, mode='lines')), use_container_width=True)

def create_radar_chart(stats):
    df = pd.DataFrame(dict(
        r=list(stats.values()),
        theta=list(stats.keys())
    ))
    fig = px.line_polar(df, r='r', theta='theta', line_close=True)
    fig.update_traces(fill='toself')
    fig.update_layout(
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 10])
        ),
        showlegend=False
    )
    return fig

def main():
    st.set_page_config(page_title="AUDSOL - AI Automation for Writers", page_icon="📚", layout="wide")

    menu = ["Home", "Create Character", "Chat with Characters", "Generate Book", "Convert to Audiobook", 
            "Book Outline Generator", "Character Development Workshop", "Writing Prompts Generator", "World-Building Assistant", "Interactive Character Board"]
    choice = st.sidebar.selectbox("Menu", menu)

    if choice == "Home":
        st.title("AUDSOL - AI Automation for Writers")
        st.write("Welcome to AUDSOL, your AI-powered writing assistant!")
        
        st.header("Unlock Your Writing Potential with AI")
        st.write("""
        AUDSOL is designed to revolutionize your writing process and help you generate thousands of dollars per month through KDP self-publishing. Our AI-powered tools streamline your workflow, boost creativity, and enhance your productivity.
        """)
        
        st.subheader("Key Features:")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            - **AI Character Creation**: Bring your characters to life with unique voices and personalities.
            - **Interactive Character Chats**: Develop your characters through dynamic conversations.
            - **Automated Book Generation**: Create full-length books with AI assistance.
            - **Audiobook Conversion**: Transform your text into professional audiobooks.
            """)
        
        with col2:
            st.markdown("""
            - **Book Outline Generator**: Craft detailed outlines for your next bestseller.
            - **Character Development Workshop**: Deep dive into character creation and evolution.
            - **Writing Prompts Generator**: Spark your creativity with AI-generated prompts.
            - **Multi-format Export**: Save your work in TXT, DOCX, and PDF formats.
            """)
        
        st.header("How AUDSOL Boosts Your KDP Earnings")
        st.write("""
        1. **Rapid Content Creation**: Generate high-quality books faster than ever before.
        2. **Diverse Genre Expertise**: Our AI adapts to any genre, helping you tap into lucrative markets.
        3. **Consistent Output**: Maintain a steady publishing schedule to build your author brand.
        4. **Enhanced Quality**: AI-assisted editing and character development improve your storytelling.
        5. **Audiobook Integration**: Easily create audiobooks to diversify your income streams.
        """)
        
        st.header("Get Started Today!")
        st.write("""
        Explore our features using the sidebar menu and start transforming your writing career. With AUDSOL, you're not just writing—you're crafting your path to financial success in the world of self-publishing.
        """)
    
    elif choice == "Create Character":
        st.header("Create a New Character")
        col1, col2 = st.columns(2)
        
        with col1:
            new_name = st.text_input("Character Name")
            new_description = st.text_area("Short Character Description (500 chars)", max_chars=500)
            new_voice_file = st.file_uploader("Upload Voice Sample (MP3)", type=["mp3"])
        
        with col2:
            new_detailed_description = st.text_area("Detailed Character Description")
            new_image_file = st.file_uploader("Upload Character Image (Optional)", type=["jpg", "png"])
        
        if st.button("Create Character") and new_name and new_description and new_voice_file:
            with st.spinner("Creating character..."):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                    tmp_file.write(new_voice_file.getvalue())
                    tmp_file_path = tmp_file.name
                create_character(new_name, new_description, tmp_file_path, new_image_file)
                os.unlink(tmp_file_path)
                st.success(f"Character '{new_name}' created successfully!")
    
    elif choice == "Chat with Characters":
        st.header("Chat with Characters")
    
        if 'messages' not in st.session_state:
            st.session_state.messages = []
    
        if 'audio_transcription' not in st.session_state:
            st.session_state.audio_transcription = None
    
        if 'voice_input_stage' not in st.session_state:
            st.session_state.voice_input_stage = 'ready'
    
        if 'audio_bytes' not in st.session_state:
            st.session_state.audio_bytes = None
    
        col1, col2 = st.columns([1, 3])
    
        with col1:
            st.subheader("Characters")
            for character in st.session_state.characters:
                if st.button(character, key=f"btn_{character}"):
                    st.session_state.current_character = character
                    st.session_state.messages = []
                    st.session_state.audio_transcription = None
                    st.session_state.voice_input_stage = 'ready'
                    st.session_state.audio_bytes = None
    
        with col2:
            if 'current_character' in st.session_state and st.session_state.current_character in st.session_state.characters:
                current_char = st.session_state.characters[st.session_state.current_character]
                st.subheader(f"Chatting with {st.session_state.current_character}")
                if "image" in current_char and current_char["image"]:
                    st.image(current_char["image"], width=200)
    
                # Display chat history
                for message in st.session_state.messages:
                    with st.chat_message(message["role"]):
                        st.write(message["content"])
                        if "audio" in message:
                            st.audio(message["audio"])
    
                language = st.selectbox("Select Language", ["English", "Spanish", "French", "German", "Chinese", "Japanese"])
                input_method = st.radio("Choose input method:", ("Text", "Voice"))
    
                user_input = None
                if input_method == "Text":
                    user_input = st.chat_input("Type your message here...")
                else:
                    st.write("Voice Input Process:")
                    
                    if st.session_state.voice_input_stage == 'ready':
                        if st.button("Start Recording"):
                            st.session_state.voice_input_stage = 'recording'
                            st.experimental_rerun()
    
                    elif st.session_state.voice_input_stage == 'recording':
                        st.write("Recording... Click 'Stop Recording' when finished.")
                        audio_bytes = audio_recorder(key="voice_recorder", icon_size="2x")
                        if audio_bytes:
                            st.session_state.audio_bytes = audio_bytes
                        if st.button("Stop Recording"):
                            if st.session_state.audio_bytes:
                                st.session_state.voice_input_stage = 'transcribing'
                                st.experimental_rerun()
                            else:
                                st.warning("No audio recorded. Please try again.")
                                st.session_state.voice_input_stage = 'ready'
    
                    elif st.session_state.voice_input_stage == 'transcribing':
                        st.write("Transcribing audio...")
                        with st.spinner("Processing..."):
                            if st.session_state.audio_bytes:
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                                    tmp_file.write(st.session_state.audio_bytes)
                                    tmp_file_path = tmp_file.name
    
                                recognizer = sr.Recognizer()
                                with sr.AudioFile(tmp_file_path) as source:
                                    audio_data = recognizer.record(source)
    
                                try:
                                    transcription = recognizer.recognize_google(audio_data)
                                    st.session_state.audio_transcription = transcription
                                    user_input = transcription
                                except sr.UnknownValueError:
                                    st.warning("Speech recognition could not understand the audio. Please try again.")
                                    st.session_state.voice_input_stage = 'ready'
                                except sr.RequestError as e:
                                    st.error(f"Could not request results from speech recognition service. Please check your internet connection and try again. Error: {e}")
                                    st.session_state.voice_input_stage = 'ready'
    
                                os.unlink(tmp_file_path)
                            else:
                                st.error("No audio data available. Please record your message again.")
                                st.session_state.voice_input_stage = 'ready'
    
                if user_input:
                    # Add user message to chat history
                    st.session_state.messages.append({"role": "user", "content": user_input})
                    
                    # Display user message
                    with st.chat_message("user"):
                        st.write(user_input)
                    
                    with st.spinner("Generating response..."):
                        ai_response = generate_ai_response(st.session_state.current_character, user_input, language)
                        audio = text_to_speech(ai_response, current_char["voice_id"])
    
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                        save(audio, tmp_file.name)
                        audio_path = tmp_file.name
    
                    # Add AI response to chat history
                    st.session_state.messages.append({
                        "role": "assistant", 
                        "content": ai_response, 
                        "audio": audio_path
                    })
    
                    # Display AI response
                    with st.chat_message("assistant"):
                        st.write(ai_response)
                        st.audio(audio_path)
    
                    # Reset for next interaction
                    st.session_state.audio_transcription = None
                    st.session_state.audio_bytes = None
                    st.session_state.voice_input_stage = 'ready'
                    st.experimental_rerun()
    
            else:
                st.info("Please select a character to start chatting.")
    
    elif choice == "Generate Book":
        st.header("Generate a Book")
        
        genre_type = st.radio("Choose genre type:", ("Fiction", "Non-fiction"))
        
        if genre_type == "Fiction":
            genre = st.selectbox("Select Fiction Genre", fiction_genres)
        else:
            genre = st.selectbox("Select Non-fiction Genre", non_fiction_genres)
        
        title = st.text_input("Book Title")
        author = st.text_input("Author Name")
        description = st.text_area("Book Description")
        
        num_chapters = st.number_input("Number of Chapters", min_value=1, value=5)
        
        chapters = []
        for i in range(num_chapters):
            st.subheader(f"Chapter {i+1}")
            chapter_name = st.text_input(f"Chapter {i+1} Name")
            if genre_type == "Fiction":
                num_scenes = st.number_input(f"Number of Scenes in Chapter {i+1}", min_value=1, value=3)
                scenes = []
                for j in range(num_scenes):
                    scene_description = st.text_area(f"Scene {j+1} Description (Chapter {i+1})")
                    scenes.append(scene_description)
                chapters.append({"name": chapter_name, "scenes": scenes})
            else:
                chapter_description = st.text_area(f"Chapter {i+1} Description")
                num_parts = st.number_input(f"Number of Parts in Chapter {i+1}", min_value=1, value=3)
                parts = []
                for j in range(num_parts):
                    part_description = st.text_area(f"Part {j+1} Description (Chapter {i+1})")
                    parts.append(part_description)
                chapters.append({"name": chapter_name, "description": chapter_description, "parts": parts})
        
        if st.button("Generate Book"):
            book_details = {
                "genre_type": genre_type,
                "genre": genre,
                "title": title,
                "author": author,
                "description": description,
                "chapters": chapters
            }
            
            st.session_state.book_details = book_details
            st.session_state.generated_book = ""
            
            with st.spinner("Generating book introduction..."):
                introduction = generate_book_introduction(book_details)
                st.session_state.generated_book += f"# {title}\nBy {author}\n\n{introduction}\n\n"
            
            toc = generate_table_of_contents(book_details)
            st.session_state.generated_book += f"{toc}\n\n"
            
            for i, chapter in enumerate(chapters):
                with st.spinner(f"Generating Chapter {i+1}: {chapter['name']}..."):
                    chapter_content = generate_book_chapter(book_details, i)
                    st.session_state.generated_book += f"# Chapter {i+1}: {chapter['name']}\n\n{chapter_content}\n\n"
                st.success(f"Chapter {i+1} generated successfully!")
            
            st.success("Book generated successfully!")
            st.text_area("Generated Book", st.session_state.generated_book, height=300)
            
            save_book_formats(st.session_state.generated_book, title)
            
            st.download_button(
                label="Download as TXT",
                data=st.session_state.generated_book,
                file_name=f"{title}.txt",
                mime="text/plain"
            )
            
            with open(f"{title}.docx", "rb") as docx_file:
                st.download_button(
                    label="Download as DOCX",
                    data=docx_file,
                    file_name=f"{title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with open(f"{title}.pdf", "rb") as pdf_file:
                st.download_button(
                    label="Download as PDF",
                    data=pdf_file,
                    file_name=f"{title}.pdf",
                    mime="application/pdf"
                )

    elif choice == "Convert to Audiobook":
        st.header("Convert to Audiobook")
        st.write("Transform your written work into an engaging audiobook.")
    
        col1, col2 = st.columns([2, 1])
    
        with col1:
            upload_type = st.radio("Choose input method:", ("Upload File", "Paste Text"), horizontal=True)
    
            if upload_type == "Upload File":
                uploaded_file = st.file_uploader("Upload your book file", type=['txt', 'pdf', 'doc', 'docx'])
                if uploaded_file:
                    with st.spinner("Extracting text from file..."):
                        text = extract_text_from_file(uploaded_file)
                    st.success("Text extracted successfully!")
            else:
                text = st.text_area("Enter your book text here", height=300, placeholder="Paste your book content here...")
    
            if 'text' in locals() and text:
                st.write("Preview:")
                st.info(text[:500] + "..." if len(text) > 500 else text)
    
        with col2:
            st.subheader("Audiobook Settings")
            language = st.selectbox("Language", ["English", "Spanish", "French", "German", "Chinese", "Japanese"])
    
            voice_option = st.radio("Voice Option:", ("Predefined", "Clone Your Voice"))
    
            if voice_option == "Predefined":
                available_voices = [
                    ("21m00Tcm4TlvDq8ikWAM", "Rachel"),
                    ("2EiwWnXFnvU5JabPnv8n", "Domi"),
                    ("AZnzlk1XvdvUeBnXmlld", "Bella"),
                    ("EXAVITQu4vr4xnSDxMaL", "Antoni"),
                    ("ErXwobaYiN019PkySvjV", "Elli"),
                    ("MF3mGyEYCl7XYWbV9V6O", "Adam"),
                    ("TxGEqnHWrfWFTfGW9XjX", "Sam"),
                    ("VR6AewLTigWG4xSOukaG", "Josh"),
                    ("pNInz6obpgDQGcFmaJgB", "Arnold"),
                    ("yoZ06aMxZJJ28mfd3POQ", "Emma"),
                    ("ZQe5CZNOzWyzPSCn5a3c", "Ava"),
                    ("jBpfuIE2acCO8z3wKNLl", "Bella"),
                    ("onwK4e9ZLuTAKqWW03F9", "Dorothy"),
                    ("g5CIjZEefAph4nQFvHAz", "Josh"),
                    ("wViXBPUzp2ZZixB1xQuM", "Arnold"),
                    ("zrHiDhphv9ZnVXBqCLjz", "Charlotte"),
                    ("t0jbNlBVZ17f02VDIeMI", "Matilda"),
                    ("piTKgcLEGmPE4e6mEKli", "Matthew"),
                    ("RErgWrLnpU47n1wBUr3G", "James"),
                    ("N2lVS1w4EtoT3dr4eOWO", "Joseph"),
                    ("ODq5zmih8GrVes37Dizd", "Jeremy"),
                    ("SOYHLrjzK2X1ezoPC6cr", "Michael"),
                    ("TX3LPaxmHKxFdv7VOQHJ", "Robert"),
                    ("flq6f7yk4E4fJM5XTYNq", "Olivia"),
                    ("z9fAnlkpzviPz146aGWa", "Sophia"),
                    ("3KehXGbKd6Gu5T9fjnVK", "Ethan"),
                    ("IKne3meq5aSn9XLyUdCD", "Charlie"),
                    ("LcfcDJNUP1GQjkzn1xUU", "Amelia"),
                    ("G4i3RVLNXz6PiE5QAgCu", "Lily"),
                    ("XB0fDUnXU5powFXDhCwa", "Jackson"),
                    ("xVe4PmtsSGVPkKo5V1gB", "Chloe"),
                    ("YD5YTsM3TLX99VqWelFo", "Madison"),
                    ("jsCqWAovK2LkecY7zXl4", "Ryan"),
                    ("oWAxZDx7w5VEj9dCyTzz", "Luna"),
                    ("1vXVqeHAitMxEfPe88RV", "Maverick"),
                    ("2eGXxRfOxHPE1ZPCIzNv", "Zoe"),
                    ("tL2Ybu6R2XEfqotJ1UWp", "Thomas"),
                    ("XrExE9yKIg1WjnnlVkGX", "Scarlett"),
                    ("Yko7PKHZNXotIFUBG7I9", "Daniel"),
                    ("R2wc5f0yV8dVe4Y24eDC", "Hannah"),
                    ("qNAGtHKepPLQO6GXVImS", "Liam"),
                    ("h0ZXKKEOGPMlAb7yFQf0", "Isabella"),
                    ("Jb5FM5xWS8mhCk0DHjmx", "Grace"),
                    ("pOnD5fLH3IhEBY5YOHP4", "Noah"),
                    ("r5U4RGJuA4kh2Umi8P3J", "Oliver"),
                    ("x8GtTQIVJ5dxmYN4vffj", "Emily"),
                    ("j9FdgHYQX3nG7qjc9VVG", "Sophie"),
                    ("C8lN8YDucvTvl6yKiokH", "William"),
                    ("JM6niiVGNJTotuzPD5kd", "Evelyn"),
                    ("7JrRqOwEPGAznN7Uf8hU", "Benjamin")
                ]
                
                selected_voice = st.selectbox(
                    "Select a voice",
                    options=available_voices,
                    format_func=lambda x: f"{x[1]} (ID: {x[0]})"
                )
                
                if selected_voice:
                    voice_id, voice_name = selected_voice
                    st.write(f"Selected Voice: {voice_name}")
                    st.write(f"Voice ID: {voice_id}")
            else:
                voice_sample = st.file_uploader("Upload your voice sample (MP3)", type=["mp3"])
                if voice_sample:
                    with st.spinner("Cloning voice..."):
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                            tmp_file.write(voice_sample.getvalue())
                            tmp_file_path = tmp_file.name
                        
                        cloned_voice = elevenlabs_client.clone(
                            name="Custom Voice",
                            description="Custom cloned voice",
                            files=[tmp_file_path]
                        )
                        # No need to unlink here; we will do it later
                    voice_id = cloned_voice.voice_id
                    st.success("Voice cloned successfully!")
    
            speed = st.slider("Narration Speed", min_value=0.5, max_value=2.0, value=1.0, step=0.1)
    
        if st.button("Convert to Audiobook", type="primary"):
            if 'text' in locals() and text and 'voice_id' in locals():
                chunks = split_text(text)
                total_chunks = len(chunks)

                headers = {
                    "Accept": "audio/mpeg",
                    "Content-Type": "application/json",
                    "xi-api-key": XI_API_KEY
                }

                all_audio_data = BytesIO()

                progress_bar = st.progress(0)
                status_text = st.empty()

                for i, chunk in enumerate(chunks, 1):
                    status_text.text(f"Converting chunk {i}/{total_chunks}...")
                    try:
                        audio_content = generate_audio_chunk(chunk, voice_id, headers)
                        all_audio_data.write(audio_content)
                        progress_bar.progress(i / total_chunks)
                    except Exception as e:
                        st.error(f"Error processing chunk {i}: {str(e)}")
                        break
                    
                if all_audio_data.tell() > 0:
                    all_audio_data.seek(0)
                    st.success("Audiobook created successfully!")
                    st.audio(all_audio_data, format='audio/mp3')

                    st.download_button(
                        label="Download Complete Audiobook",
                        data=all_audio_data,
                        file_name="complete_audiobook.mp3",
                        mime="audio/mpeg",
                        key="download_complete_audiobook"
                    )
                else:
                    st.error("Failed to generate the complete audiobook.")
            else:
                st.warning("Please provide text and select a voice before converting.")

    # New feature: Book Outline Generator
    elif choice == "Book Outline Generator":
        st.header("Detailed Book Outline Generator")
        st.write("Create a comprehensive, chapter-by-chapter outline for your next bestseller.")

        col1, col2 = st.columns([2, 1])

        with col1:
            title = st.text_input("Book Title", placeholder="Enter your book title")
            main_idea = st.text_area("Main Idea or Concept", placeholder="Describe the central theme or concept of your book in detail", height=150)
            target_audience = st.text_input("Target Audience", placeholder="Who is your book for? Be specific about demographics, interests, etc.")
            key_themes = st.text_area("Key Themes or Topics", placeholder="List the main themes or topics you want to cover in your book", height=100)

        with col2:
            genre_type = st.radio("Genre Type:", ("Fiction", "Non-fiction"))
            if genre_type == "Fiction":
                genre = st.selectbox("Fiction Genre", fiction_genres)
                protagonist = st.text_input("Protagonist", placeholder="Describe your main character")
                setting = st.text_input("Setting", placeholder="Where and when does your story take place?")
            else:
                genre = st.selectbox("Non-fiction Genre", non_fiction_genres)
                expertise_level = st.select_slider("Reader's Expertise Level", options=["Beginner", "Intermediate", "Advanced", "Expert"])

            desired_length = st.number_input("Estimated Word Count", min_value=10000, value=80000, step=5000, help="Approximate length of your book")
            num_chapters = st.number_input("Number of Chapters", min_value=5, max_value=50, value=15)

        if st.button("Generate Detailed Outline", type="primary"):
            if title and main_idea and target_audience and genre and key_themes:
                with st.spinner("Crafting your comprehensive book outline..."):
                    if genre_type == "Fiction":
                        outline_prompt = f"""Generate an extremely detailed book outline for a {genre} novel titled '{title}'. 
                        Main idea: {main_idea}
                        Target audience: {target_audience}
                        Key themes: {key_themes}
                        Protagonist: {protagonist}
                        Setting: {setting}
                        Estimated length: {desired_length} words
                        Number of chapters: {num_chapters}

                        For each chapter, provide:
                        1. A compelling chapter title
                        2. A detailed synopsis (500-800 words)
                        3. Key plot points or events
                        4. Character development and interactions
                        5. Setting details and atmosphere
                        6. Themes explored in the chapter
                        7. Any foreshadowing or plot twists
                        8. Estimated word count for the chapter

                        Additionally, include:
                        - An engaging prologue idea
                        - A captivating epilogue concept
                        - Suggestions for potential subplots
                        - Ideas for symbolic elements or motifs throughout the book

                        Make the outline as comprehensive and detailed as possible, using the maximum available tokens."""
                    else:
                        outline_prompt = f"""Generate an extremely detailed book outline for a {genre} non-fiction book titled '{title}'. 
                        Main idea: {main_idea}
                        Target audience: {target_audience} (Expertise level: {expertise_level})
                        Key themes or topics: {key_themes}
                        Estimated length: {desired_length} words
                        Number of chapters: {num_chapters}

                        For each chapter, provide:
                        1. An informative chapter title
                        2. A detailed chapter summary (200-300 words)
                        3. Main concepts or arguments presented
                        4. Supporting evidence, data, or examples to include
                        5. Potential expert quotes or case studies to research
                        6. Practical applications or exercises for readers
                        7. Key takeaways from the chapter
                        8. Estimated word count for the chapter

                        Additionally, include:
                        - An attention-grabbing introduction outline
                        - A powerful conclusion and call-to-action outline
                        - Ideas for sidebars, infographics, or illustrations
                        - Suggestions for further reading or resources

                        Make the outline as comprehensive and detailed as possible, using the maximum available tokens."""

                    outline = together_client.chat.completions.create(
                        model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
                        messages=[
                            {"role": "system", "content": "You are a professional book outliner and developmental editor with extensive experience in creating detailed, chapter-by-chapter outlines for bestselling books across various genres."},
                            {"role": "user", "content": outline_prompt}
                        ],
                        max_tokens=15000,
                        temperature=0.7
                    )

                    st.session_state.book_outline = outline.choices[0].message.content

                st.success("Comprehensive book outline generated successfully!")

                # Display the outline in a structured and visually appealing way
                st.subheader("Your Detailed Book Outline")
                outline_lines = st.session_state.book_outline.split('\n')
                chapter_count = 0
                for line in outline_lines:
                    if line.strip().startswith('Chapter'):
                        chapter_count += 1
                        st.markdown(f"<h3 style='color: #1e90ff;'>{line.strip()}</h3>", unsafe_allow_html=True)
                    elif any(section in line for section in ['Synopsis:', 'Summary:', 'Plot Points:', 'Main Concepts:', 'Character Development:', 'Supporting Evidence:', 'Setting:', 'Practical Applications:', 'Themes:', 'Key Takeaways:', 'Foreshadowing:', 'Estimated Word Count:']):
                        st.markdown(f"<h4 style='color: #32cd32;'>{line.strip()}</h4>", unsafe_allow_html=True)
                    else:
                        st.write(line.strip())

                st.info(f"Total Chapters: {chapter_count}")

                st.download_button(
                    label="Download Detailed Outline",
                    data=st.session_state.book_outline,
                    file_name=f"{title.replace(' ', '_').lower()}_detailed_outline.txt",
                    mime="text/plain",
                    key="download_detailed_outline"
                )
            else:
                st.warning("Please fill in all the required fields to generate a detailed outline.")

    # New feature: Character Development Workshop
    elif choice == "Character Development Workshop":
        st.header("Character Development Workshop")
        st.write("Bring your characters to life with our in-depth development tools.")

        col1, col2 = st.columns([1, 1])

        with col1:
            character_name = st.text_input("Character Name", placeholder="Enter character's name")
            character_role = st.selectbox("Character Role", ["Protagonist", "Antagonist", "Supporting Character", "Mentor", "Love Interest", "Sidekick"])
            character_age = st.number_input("Age", min_value=0, max_value=250, value=30)
            character_occupation = st.text_input("Occupation", placeholder="Character's job or main activity")

        with col2:
            character_background = st.text_area("Background", placeholder="Brief history or backstory", height=100)
            character_goals = st.text_area("Goals", placeholder="What does the character want to achieve?", height=80)
            character_fears = st.text_area("Fears or Weaknesses", placeholder="What holds the character back?", height=80)

        col3, col4 = st.columns([1, 1])

        with col3:
            physical_attributes = st.text_area("Physical Attributes", placeholder="Describe appearance, mannerisms, etc.", height=100)

        with col4:
            personality_traits = st.text_area("Personality Traits", placeholder="List key personality characteristics", height=100)

        if st.button("Develop Character", type="primary"):
            if character_name and character_role and character_background:
                with st.spinner("Crafting your character profile..."):
                    character_prompt = f"""
                    Create a detailed character profile for {character_name}:
                    - Age: {character_age}
                    - Occupation: {character_occupation}
                    - Role: {character_role}
                    - Background: {character_background}
                    - Goals: {character_goals}
                    - Fears/Weaknesses: {character_fears}
                    - Physical Attributes: {physical_attributes}
                    - Personality Traits: {personality_traits}

                    Expand on these details to create a rich, multi-dimensional character. Include potential character arc, quirks, and how they might interact with other characters or drive the plot forward.
                    """

                    character_profile = together_client.chat.completions.create(
                        model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
                        messages=[
                            {"role": "system", "content": "You are a professional character developer for novels and screenplays, skilled in creating complex, believable characters."},
                            {"role": "user", "content": character_prompt}
                        ],
                        max_tokens=6500,
                        temperature=0.7
                    )

                    st.session_state.character_profile = character_profile.choices[0].message.content

                st.success("Character profile developed successfully!")

                # Display the character profile in a more structured way
                st.subheader(f"{character_name}'s Character Profile")
                profile_lines = st.session_state.character_profile.split('\n')
                for line in profile_lines:
                    if ':' in line:
                        key, value = line.split(':', 1)
                        st.markdown(f"**{key.strip()}:** {value.strip()}")
                    else:
                        st.write(line.strip())

                st.download_button(
                    label="Download Character Profile",
                    data=st.session_state.character_profile,
                    file_name=f"{character_name.replace(' ', '_').lower()}_profile.txt",
                    mime="text/plain",
                    key="download_character_profile"
                )
            else:
                st.warning("Please fill in at least the character's name, role, and background to generate a profile.")

    # New feature: Writing Prompts Generator
    elif choice == "Writing Prompts Generator":
        st.header("Writing Prompts Generator")
        st.write("Spark your creativity with custom writing prompts.")

        col1, col2 = st.columns([1, 1])

        with col1:
            prompt_type = st.selectbox("Prompt Type", ["General", "Sci-Fi", "Fantasy", "Romance", "Mystery", "Historical", "Horror", "Thriller"])
            prompt_length = st.slider("Prompt Complexity", min_value=1, max_value=5, value=3, help="1: Simple, 5: Elaborate")

        with col2:
            specific_elements = st.multiselect("Include Specific Elements", ["Character", "Setting", "Conflict", "Theme", "Plot Twist"])
            writing_style = st.selectbox("Writing Style", ["Any", "Descriptive", "Dialogue-heavy", "Action-packed", "Introspective", "Humorous"])

        mood = st.select_slider("Mood", options=["Dark", "Neutral", "Light"], value="Neutral")

        if st.button("Generate Writing Prompt", type="primary"):
            with st.spinner("Crafting your writing prompt..."):
                prompt_request = f"""
                Generate a {prompt_type.lower()} writing prompt.
                Complexity: {prompt_length}/5
                Include these elements: {', '.join(specific_elements)}
                Writing style: {writing_style}
                Mood: {mood}

                The prompt should inspire a short story or scene that a writer can immediately start working on.
                """

                generated_prompt = together_client.chat.completions.create(
                    model="meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo",
                    messages=[
                        {"role": "system", "content": "You are a creative writing prompt generator, skilled in crafting inspiring and thought-provoking prompts for writers."},
                        {"role": "user", "content": prompt_request}
                    ],
                    max_tokens=1024,
                    temperature=0.9
                )

                st.session_state.writing_prompt = generated_prompt.choices[0].message.content

            st.success("Writing prompt generated successfully!")

            # Display the writing prompt in an attractive format
            st.markdown("### Your Writing Prompt")
            st.info(st.session_state.writing_prompt)

            col3, col4 = st.columns([1, 1])
            with col3:
                if st.button("Generate Another Prompt"):
                    st.experimental_rerun()
            with col4:
                st.download_button(
                    label="Save Prompt",
                    data=st.session_state.writing_prompt,
                    file_name="writing_prompt.txt",
                    mime="text/plain",
                    key="save_writing_prompt"
                )

    elif choice == "Interactive Character Board":
        st.header("Interactive Character Board")
    
        # Initialize session state variables
        if "character_board" not in st.session_state:
            st.session_state.character_board = []
        if "custom_stats" not in st.session_state:
            st.session_state.custom_stats = ["Intelligence", "Strength", "Speed", "Durability", "Energy Projection", "Fighting Skills"]

        # Character creation form
        st.subheader("Add New Character")
        with st.form("character_form"):
            new_name = st.text_input("Character Name")
            new_description = st.text_area("Character Description", max_chars=1000)
            new_image = st.file_uploader("Character Image", type=["jpg", "png", "jpeg"])

            st.subheader("Character Stats")
            stats = {}
            for stat in st.session_state.custom_stats:
                stats[stat] = st.slider(f"{stat}", 0, 10, 5)

            background = st.text_area("Character Background")
            abilities = st.text_area("Special Abilities")
            weaknesses = st.text_area("Weaknesses")
            relationships = st.text_area("Relationships")
            additional_info = st.text_area("Additional Information")

            submit_button = st.form_submit_button("Add Character")

        if submit_button:
            if new_name and new_description:
                character_data = {
                    "name": new_name,
                    "description": new_description,
                    "stats": stats,
                    "background": background,
                    "abilities": abilities,
                    "weaknesses": weaknesses,
                    "relationships": relationships,
                    "additional_info": additional_info,
                }
        
                if new_image:
                    image = Image.open(new_image)
                    img_byte_arr = io.BytesIO()
                    image.save(img_byte_arr, format='PNG')
                    character_data["image"] = img_byte_arr.getvalue()
        
                st.session_state.character_board.append(character_data)
        
                st.success(f"Character '{new_name}' added successfully!")
            else:
                st.warning("Please provide at least a name and description for the character.")
        
        # Stat management
        with st.expander("Manage Stats"):
            st.subheader("Current Stats")
            for stat in st.session_state.custom_stats:
                col1, col2 = st.columns([3, 1])
                col1.write(stat)
                if col2.button("Remove", key=f"remove_{stat}"):
                    st.session_state.custom_stats.remove(stat)
                    st.experimental_rerun()
        
            new_stat = st.text_input("Add New Stat")
            if st.button("Add Stat"):
                if new_stat and new_stat not in st.session_state.custom_stats:
                    st.session_state.custom_stats.append(new_stat)
                    st.experimental_rerun()
                else:
                    st.warning("Stat already exists or is empty.")

        # Display character cards
        st.subheader("Character Cards")
        if st.session_state.character_board:
            for idx, character in enumerate(st.session_state.character_board):
                with st.expander(character["name"], expanded=True):
                    col1, col2 = st.columns([1, 1])
                    with col1:
                        if "image" in character:
                            st.image(character["image"], use_column_width=True)
                        st.write(f"**Description:** {character['description']}")
                        st.write(f"**Background:** {character['background']}")
                        st.write(f"**Special Abilities:** {character['abilities']}")
                        st.write(f"**Weaknesses:** {character['weaknesses']}")
                        st.write(f"**Relationships:** {character['relationships']}")
                        if character["additional_info"]:
                            st.write(f"**Additional Info:** {character['additional_info']}")
                    with col2:
                        st.subheader("Character Stats")
                        fig = create_radar_chart(character["stats"])
                        st.plotly_chart(fig, use_container_width=True)

                    if st.button(f"Delete {character['name']}", key=f"delete_{idx}"):
                        st.session_state.character_board.pop(idx)
                        st.experimental_rerun()
        else:
            st.info("No characters added yet. Use the form above to add characters to the board.")

    elif choice == "World-Building Assistant":
        world_building_assistant()

        # Optional: Add a feature to view saved prompts
        if st.checkbox("View Saved Prompts"):
            if "saved_prompts" not in st.session_state:
                st.session_state.saved_prompts = []

            if st.session_state.writing_prompt and st.button("Add Current Prompt to Saved"):
                st.session_state.saved_prompts.append(st.session_state.writing_prompt)
                st.success("Prompt added to saved list!")

            if st.session_state.saved_prompts:
                for i, prompt in enumerate(st.session_state.saved_prompts):
                    st.text_area(f"Saved Prompt {i+1}", prompt, height=100, key=f"saved_prompt_{i}")
            else:
                st.info("No saved prompts yet. Generate and save some prompts to see them here!")

    # Sidebar for quick access to generated content
    with st.sidebar:
        st.header("Quick Access")
        if "generated_book" in st.session_state:
            if st.button("View Generated Book"):
                st.session_state.current_view = "generated_book"
        if "book_outline" in st.session_state:
            if st.button("View Book Outline"):
                st.session_state.current_view = "book_outline"
        if "character_profile" in st.session_state:
            if st.button("View Character Profile"):
                st.session_state.current_view = "character_profile"
        if "writing_prompt" in st.session_state:
            if st.button("View Writing Prompt"):
                st.session_state.current_view = "writing_prompt"
        if "story_branches" in st.session_state:
            if st.button("View Interactive Story"):
                st.session_state.current_view = "interactive_story"
        if "world_elements" in st.session_state:
            if st.button("View World-Building"):
                st.session_state.current_view = "world_building"

    # Display the selected content in the main area
    if "current_view" in st.session_state:
        if st.session_state.current_view == "generated_book":
            st.header("Generated Book")
            st.text_area("Book Content", st.session_state.generated_book, height=400)
        elif st.session_state.current_view == "book_outline":
            st.header("Book Outline")
            st.text_area("Outline", st.session_state.book_outline, height=400)
        elif st.session_state.current_view == "character_profile":
            st.header("Character Profile")
            st.text_area("Profile", st.session_state.character_profile, height=400)
        elif st.session_state.current_view == "writing_prompt":
            st.header("Writing Prompt")
            st.text_area("Prompt", st.session_state.writing_prompt, height=200)
        elif st.session_state.current_view == "world_building":
            world_building_assistant()
if __name__ == "__main__":
    main()
